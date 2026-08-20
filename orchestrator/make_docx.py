#!/usr/bin/env python3
"""
make_docx.py

Two-part build for the fourth draft of "The System of You":

PART 1 - Clean the chapter markdown files (ch00.md .. ch13.md) in output/draft4/:
  * Remove leaked working-stage label lines near the top of each file.
  * Strip any existing top heading / leading separators, then prepend a
    canonical `# <heading>` line.
  * Special case: ch10.md is a broken scrubber artifact (only an "Artifact
    Removal Log"), so the real chapter body is sourced from ch10_raw.md.
    Per author decision, the trailing "Rewrite Log" section in ch10_raw.md
    is KEPT as part of the chapter content.

PART 2 - Build output/full_fourth_draft.docx with python-docx, plus regenerate
  output/full_fourth_draft.md by concatenating the cleaned chapters.

Run: python3 orchestrator/make_docx.py
"""

import os
import re

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)                       # project root: .../Charlotte
DRAFT_DIR = os.path.join(ROOT, "output", "draft4")
OUTPUT_DIR = os.path.join(ROOT, "output")
DOCX_PATH = os.path.join(OUTPUT_DIR, "full_fourth_draft.docx")
MD_PATH = os.path.join(OUTPUT_DIR, "full_fourth_draft.md")

BOOK_TITLE = "THE SYSTEM OF YOU"
BOOK_SUBTITLE = "The Manual You Were Never Given"

# Canonical headings, in book order (chapter-id -> heading text)
CANONICAL = [
    ("ch00", "A Note Before We Start"),
    ("ch01", "Chapter 1: The Doorway"),
    ("ch02", "Chapter 2: Body"),
    ("ch03", "Chapter 3: Wiring"),
    ("ch04", "Chapter 4: Habit"),
    ("ch05", "Chapter 5: Room"),
    ("ch06", "Chapter 6: Story"),
    ("ch07", "Chapter 7: The Method Made Whole"),
    ("ch08", "Chapter 8: The Toolkit"),
    ("ch09", "Chapter 9: The Hard Cases"),
    ("ch10", "Chapter 10: When the Layer Won't Move"),
    ("ch11", "Chapter 11: Belief, Then Faith"),
    ("ch12", "Chapter 12: Watching It Change \u2014 The Growth Map"),
    ("ch13", "Chapter 13: A Note Before You Go"),
]

# Chapters whose canonical body is sourced from an alternate file.
# ch10.md is broken (only a scrubber log), real body lives in ch10_raw.md.
SOURCE_OVERRIDE = {
    "ch10": "ch10_raw.md",
}

# Leaked pipeline artifacts embedded DEEP in chapter bodies that should be
# removed (per author decision). These are whole foreign documents / logs that
# leaked in during the rewrite pipeline and would otherwise render as phantom
# top-level chapters in the DOCX. Keyed by chapter id -> list of top-level
# heading texts that begin a leaked section. Everything from the first matching
# `# <heading>` line to end-of-file is removed.
#
# NOTE: This intentionally does NOT touch ch10's trailing "Rewrite Log"
# (author asked to keep that), nor ch04's "### Rewrite Log" (a sub-heading, not
# an H1). Only the specific leaked H1 documents below are stripped.
BODY_ARTIFACT_H1 = {
    "ch04": ["Model Bible"],   # the entire Model Bible reference doc leaked in
    "ch05": ["Evidence Log"],  # evidence-agent working-artifact table leaked in
}


def strip_body_artifacts(chapter_id, text):
    """
    Remove leaked H1 artifact sections embedded deep in a chapter body.
    Cuts from the first matching top-level `# <heading>` line (that is NOT the
    chapter's own canonical heading at the very top) through to end of file,
    along with any immediately preceding blank lines / horizontal rule.
    """
    markers = BODY_ARTIFACT_H1.get(chapter_id)
    if not markers:
        return text

    lines = text.split("\n")
    cut_at = None
    for i, line in enumerate(lines):
        s = line.strip()
        if i == 0:
            # never treat the canonical top heading as an artifact
            continue
        if s.startswith("#") and not s.startswith("##"):
            htext = s.lstrip("#").strip()
            # match if the heading text starts with any configured marker
            for marker in markers:
                if htext.lower().startswith(marker.lower()):
                    cut_at = i
                    break
        if cut_at is not None:
            break

    if cut_at is None:
        return text

    # walk back over any preceding blank lines / horizontal rules
    j = cut_at - 1
    while j >= 0 and (lines[j].strip() == "" or lines[j].strip() == "---"
                      or (set(lines[j].strip()) == {"-"} and len(lines[j].strip()) >= 3)):
        j -= 1

    kept = lines[:j + 1]
    # trim trailing blank lines
    while kept and kept[-1].strip() == "":
        kept.pop()
    return "\n".join(kept) + "\n"


# Trailing pipeline logs that leaked into the reader-facing text and must be cut
# from EVERY chapter. These are the working artifacts emitted by the rewrite /
# evidence / scrubber agents (Rewrite Log, Evidence Log, Artifact Removal Log,
# etc.). They always appear at the END of a chapter, so we cut from the first
# matching heading line to end-of-file, plus any preceding separator/blank lines.
_TRAILING_LOG_RE = re.compile(
    r'^\s*(?:#{1,4}\s*|\*\*)?'
    r'(Rewrite Log|Evidence Log|Artifact Removal Log|Artifact Scrub Log|'
    r'Continuity (?:Audit|Notes?)|Editorial Gate|Scrub Log)\b',
    re.IGNORECASE,
)


def strip_trailing_logs(text):
    """
    Remove any trailing pipeline log block (Rewrite Log / Evidence Log /
    Artifact Removal Log / etc.) that leaked into a chapter body. Cuts from the
    first matching log heading through end-of-file, trimming any preceding
    horizontal-rule / blank separator. Author markers ([AUTHOR VOICE] /
    [NOTE TO AUTHOR]) sit in the body above these blocks and are preserved.
    """
    lines = text.split("\n")
    cut_at = None
    for i, line in enumerate(lines):
        if _TRAILING_LOG_RE.match(line):
            cut_at = i
            break
    if cut_at is None:
        return text
    j = cut_at - 1
    while j >= 0 and _is_blank_or_sep(lines[j]):
        j -= 1
    kept = lines[:j + 1]
    while kept and kept[-1].strip() == "":
        kept.pop()
    return "\n".join(kept) + "\n"


# ---------------------------------------------------------------------------
# PART 1 helpers - label detection / cleaning
# ---------------------------------------------------------------------------

def _norm(line):
    """Lowercased text with markdown emphasis / whitespace stripped, for matching."""
    t = line.strip()
    # remove leading/trailing markdown emphasis chars and spaces
    t = t.strip("*_ ").strip()
    return t.lower()


def is_leaked_label(line):
    """Return True if this line is a leaked working-stage label that must be removed."""
    raw = line.strip()
    if not raw:
        return False

    norm = _norm(line)

    # Direct known labels
    if norm in ("clean chapter", "evidenced chapter"):
        return True

    # Any line whose stripped/lowercased text starts with these phrases
    if norm.startswith("clean chapter") or norm.startswith("evidenced chapter"):
        return True

    # Parenthetical evidenced/version notes, e.g. *(Evidenced version)*,
    # *(evidenced version)*, *(No production artifacts were detected...)*,
    # *(see full text above)*, *(Belief, Then Faith - revised ...)*,
    # *(Modifications are shown inline; ...)*
    inner = raw.strip("*_ ").strip()
    if inner.startswith("(") and inner.endswith(")"):
        low = inner.lower()
        if ("evidenced" in low or "clean chapter" in low or "revised" in low
                or "no production artifacts" in low
                or "see full text" in low
                or "modifications are shown" in low):
            return True

    return False


# Lines that are pure top-of-file separators / blanks we can safely skip while
# hunting for the heading region.
def _is_blank_or_sep(line):
    s = line.strip()
    return s == "" or s == "---" or set(s) == {"-"} and len(s) >= 3


def _is_top_heading(line):
    """A markdown heading (#...) or an all-bold title line at the very top."""
    s = line.strip()
    if s.startswith("#"):
        return True
    # entirely-bold line acting as a title, e.g. **THE DOORWAY - YOU ARE A SYSTEM**
    if s.startswith("**") and s.endswith("**") and len(s) > 4:
        return True
    return False


def clean_chapter(chapter_id, raw_text):
    """
    Clean a single chapter's text:
      * strip leaked labels from the TOP region only (first ~6 non-blank lines)
      * strip a leading top heading + leading separators/blanks
      * prepend the canonical `# <heading>`
    Deeper ##/### subheadings and all body content are preserved untouched.
    """
    heading = dict(CANONICAL)[chapter_id]
    lines = raw_text.split("\n")

    # --- Step 1: remove leaked labels found within the top region -----------
    # We scan the first non-blank lines (up to 6 counted) and drop any leaked
    # label lines. We stop scanning the "top region" once we've passed ~6
    # non-blank lines that are NOT labels/headings/separators (i.e. real body).
    cleaned = []
    nonblank_seen = 0
    top_region_done = False
    for line in lines:
        if not top_region_done:
            stripped = line.strip()
            if stripped == "":
                cleaned.append(line)
                continue
            nonblank_seen += 1
            if is_leaked_label(line):
                # drop this label line entirely
                continue
            # Once we've seen enough non-label, non-heading real content, stop.
            if nonblank_seen > 6:
                top_region_done = True
            cleaned.append(line)
        else:
            cleaned.append(line)

    # --- Step 2: strip leading blanks / separators / an existing top heading -
    idx = 0
    n = len(cleaned)
    # skip leading blanks and separators
    while idx < n and _is_blank_or_sep(cleaned[idx]):
        idx += 1
    # drop a single leading top heading (markdown # or all-bold title) if present
    if idx < n and _is_top_heading(cleaned[idx]):
        idx += 1
        # after removing the heading, also skip following blanks/separators
        while idx < n and _is_blank_or_sep(cleaned[idx]):
            idx += 1

    body_lines = cleaned[idx:]
    # trim trailing whitespace-only lines
    while body_lines and body_lines[-1].strip() == "":
        body_lines.pop()

    body = "\n".join(body_lines).rstrip()

    # --- Step 3: prepend canonical heading ----------------------------------
    result = "# " + heading + "\n\n" + body + "\n"
    return result


# ---------------------------------------------------------------------------
# PART 2 helpers - lightweight markdown -> docx converter
# ---------------------------------------------------------------------------
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

AUTHOR_RED = RGBColor(0xC0, 0x00, 0x00)

# Regex for inline spans: **bold**, *italic*, _italic_
INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|_[^_]+?_)")


def add_inline_runs(paragraph, text, base_bold=False, base_italic=False,
                    color=None):
    """Parse inline **bold**, *italic*, _italic_ spans and add runs."""
    if text == "":
        return
    pos = 0
    for m in INLINE_RE.finditer(text):
        # text before the match
        if m.start() > pos:
            _add_run(paragraph, text[pos:m.start()], base_bold, base_italic, color)
        token = m.group(0)
        if token.startswith("**") and token.endswith("**"):
            _add_run(paragraph, token[2:-2], True, base_italic, color)
        elif token.startswith("*") and token.endswith("*"):
            _add_run(paragraph, token[1:-1], base_bold, True, color)
        elif token.startswith("_") and token.endswith("_"):
            _add_run(paragraph, token[1:-1], base_bold, True, color)
        else:
            _add_run(paragraph, token, base_bold, base_italic, color)
        pos = m.end()
    if pos < len(text):
        _add_run(paragraph, text[pos:], base_bold, base_italic, color)


def _add_run(paragraph, text, bold, italic, color):
    if text == "":
        return
    run = paragraph.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if color is not None:
        run.font.color.rgb = color


def _is_author_marker(line):
    s = line.strip()
    low = s.lower()
    return ("[note to author" in low) or ("[author voice]" in low)


def _list_number_match(line):
    """Return the text after a leading '1. ' style numbered-list marker, or None."""
    m = re.match(r"^\s*\d+[.)]\s+(.*)$", line)
    if m:
        return m.group(1)
    return None


def render_markdown_to_docx(document, md_text, first_chapter):
    """Render one chapter's cleaned markdown into the docx document."""
    lines = md_text.split("\n")
    first_h1_done = False

    for line in lines:
        stripped = line.strip()

        # Skip standalone horizontal rules
        if stripped == "---" or (set(stripped) == {"-"} and len(stripped) >= 3):
            continue

        # Blank line -> paragraph separator (empty paragraph)
        if stripped == "":
            document.add_paragraph("")
            continue

        # Headings
        if stripped.startswith("####"):
            text = stripped[4:].strip()
            p = document.add_paragraph(style="Heading 4")
            add_inline_runs(p, text)
            continue
        if stripped.startswith("###"):
            text = stripped[3:].strip()
            p = document.add_paragraph(style="Heading 3")
            add_inline_runs(p, text)
            continue
        if stripped.startswith("##"):
            text = stripped[2:].strip()
            p = document.add_paragraph(style="Heading 2")
            add_inline_runs(p, text)
            continue
        if stripped.startswith("#"):
            text = stripped[1:].strip()
            # page break before each H1 chapter heading, except the very first
            if not (first_chapter and not first_h1_done):
                document.add_page_break()
            first_h1_done = True
            p = document.add_paragraph(style="Heading 1")
            add_inline_runs(p, text)
            continue

        # Author markers -> bold (red) paragraph
        if _is_author_marker(line):
            p = document.add_paragraph()
            add_inline_runs(p, stripped, base_bold=True, color=AUTHOR_RED)
            continue

        # Blockquote
        if stripped.startswith(">"):
            text = stripped[1:].strip()
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            add_inline_runs(p, text, base_italic=True)
            continue

        # Bullet list
        if stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:].strip()
            p = document.add_paragraph(style="List Bullet")
            add_inline_runs(p, text)
            continue

        # Numbered list
        num_text = _list_number_match(line)
        if num_text is not None:
            p = document.add_paragraph(style="List Number")
            add_inline_runs(p, num_text)
            continue

        # Table rows (markdown tables) -> keep as plain paragraphs so content
        # is not lost; render the row text without leading/trailing pipes.
        if stripped.startswith("|"):
            # skip separator rows like |---|---|
            cells_only = stripped.strip("|")
            if set(cells_only.replace("|", "").replace(":", "").strip()) <= {"-", " "} \
                    and "-" in cells_only:
                continue
            row_text = " | ".join(c.strip() for c in stripped.strip("|").split("|"))
            p = document.add_paragraph()
            add_inline_runs(p, row_text)
            continue

        # Whole-line bold (e.g. **SOMETHING**) -> normal paragraph, bold
        if stripped.startswith("**") and stripped.endswith("**") and \
                stripped.count("**") == 2:
            p = document.add_paragraph()
            add_inline_runs(p, stripped)  # add_inline_runs handles the ** span
            continue

        # Default: normal paragraph with inline formatting
        p = document.add_paragraph()
        add_inline_runs(p, stripped)


def build_title_page(document):
    title = document.add_paragraph(style="Title")
    trun = title.add_run(BOOK_TITLE)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle: try the "Subtitle" style, fall back to Heading 2 italic.
    try:
        sub = document.add_paragraph(style="Subtitle")
    except KeyError:
        sub = document.add_paragraph(style="Heading 2")
    srun = sub.add_run(BOOK_SUBTITLE)
    srun.italic = True
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def word_count(text):
    return len(re.findall(r"\S+", text))


def main():
    # ---- PART 1: clean each chapter and collect cleaned text --------------
    cleaned_chapters = []  # list of (chapter_id, heading, cleaned_text)
    for chapter_id, heading in CANONICAL:
        src_name = SOURCE_OVERRIDE.get(chapter_id, chapter_id + ".md")
        src_path = os.path.join(DRAFT_DIR, src_name)
        with open(src_path, "r", encoding="utf-8") as fh:
            raw = fh.read()

        cleaned = clean_chapter(chapter_id, raw)
        # Remove leaked H1 artifact sections buried deep in the body
        cleaned = strip_body_artifacts(chapter_id, cleaned)
        # Cut any trailing pipeline log block leaked into the body (all chapters)
        cleaned = strip_trailing_logs(cleaned)

        # Write cleaned content back to chXX.md (fixes broken ch10.md too)
        dest_path = os.path.join(DRAFT_DIR, chapter_id + ".md")
        with open(dest_path, "w", encoding="utf-8") as fh:
            fh.write(cleaned)

        cleaned_chapters.append((chapter_id, heading, cleaned))

    # ---- PART 2a: build the DOCX ------------------------------------------
    document = Document()
    build_title_page(document)

    for i, (chapter_id, heading, text) in enumerate(cleaned_chapters):
        first_chapter = (i == 0)
        render_markdown_to_docx(document, text, first_chapter)

    document.save(DOCX_PATH)

    # ---- PART 2b: regenerate the concatenated markdown --------------------
    md_parts = []
    md_parts.append("# " + BOOK_TITLE)
    md_parts.append("")
    md_parts.append("*" + BOOK_SUBTITLE + "*")
    md_parts.append("")
    md_parts.append("---")
    md_parts.append("")
    for i, (chapter_id, heading, text) in enumerate(cleaned_chapters):
        md_parts.append(text.rstrip())
        md_parts.append("")
        if i < len(cleaned_chapters) - 1:
            md_parts.append("---")
            md_parts.append("")
    full_md = "\n".join(md_parts).rstrip() + "\n"
    with open(MD_PATH, "w", encoding="utf-8") as fh:
        fh.write(full_md)

    # ---- Report -----------------------------------------------------------
    total_words = sum(word_count(text) for _, _, text in cleaned_chapters)
    docx_size = os.path.getsize(DOCX_PATH)
    md_size = os.path.getsize(MD_PATH)

    print("=" * 60)
    print("BUILD COMPLETE")
    print("=" * 60)
    print("Chapters processed : {}".format(len(cleaned_chapters)))
    print("Total word count   : {:,}".format(total_words))
    print("-" * 60)
    print("DOCX written : {}".format(DOCX_PATH))
    print("   size      : {:,} bytes".format(docx_size))
    print("MD written   : {}".format(MD_PATH))
    print("   size      : {:,} bytes".format(md_size))
    print("-" * 60)
    print("Per-chapter word counts:")
    for chapter_id, heading, text in cleaned_chapters:
        print("   {:<6} {:>6} words   ({})".format(
            chapter_id, word_count(text), heading))


if __name__ == "__main__":
    main()
