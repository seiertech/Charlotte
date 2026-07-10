"""Export a Charlotte manuscript (.md) to a clean DOCX and a PDF — the .md is never modified.

The DOCX is built natively (real Word heading styles, title page, chapter page breaks,
proper tables) so it imports cleanly into Google Docs (File > Open > upload the .docx).

Usage:
    python3 orchestrator/export_formats.py output/full_second_draft.md
    # -> output/full_second_draft.docx  +  output/full_second_draft.pdf
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Pt, RGBColor


# ----------------------------------------------------------------------
# Inline formatting: **bold**, *italic*, `code`
# ----------------------------------------------------------------------
_INLINE = re.compile(r'(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)')


def add_runs(paragraph, text: str) -> None:
    for part in _INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2]); r.bold = True
        elif part.startswith("*") and part.endswith("*"):
            r = paragraph.add_run(part[1:-1]); r.italic = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1]); r.font.name = "Consolas"
        else:
            paragraph.add_run(part)


def strip_inline(text: str) -> str:
    return re.sub(r'[*`]', '', text)


# ----------------------------------------------------------------------
# DOCX builder
# ----------------------------------------------------------------------
def build_docx(md_text: str, out_path: Path) -> None:
    doc = docx.Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Georgia"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.4

    lines = md_text.splitlines()
    i = 0
    first_h1_seen = False
    n = len(lines)

    def next_content_line(idx: int) -> str:
        j = idx + 1
        while j < n and not lines[j].strip():
            j += 1
        return lines[j].strip() if j < n else ""

    while i < n:
        raw = lines[i]
        line = raw.strip()

        # blank
        if not line:
            i += 1
            continue

        # table block
        if line.startswith("|") and "|" in line[1:]:
            block = []
            while i < n and lines[i].strip().startswith("|"):
                block.append(lines[i].strip())
                i += 1
            _add_table(doc, block)
            continue

        # horizontal rule / scene break
        if re.fullmatch(r'-{3,}', line) or re.fullmatch(r'\*{3,}', line):
            # skip if it directly precedes a chapter heading (page break handles it)
            if next_content_line(i).startswith("# "):
                i += 1
                continue
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run("* * *")
            r.italic = True
            r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            i += 1
            continue

        # headings
        m = re.match(r'(#{1,6})\s+(.*)', line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            if level == 1 and not first_h1_seen:
                # Title page
                first_h1_seen = True
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(strip_inline(text)); r.bold = True; r.font.size = Pt(26)
            elif level == 1:
                h = doc.add_heading(strip_inline(text), level=1)
                h.paragraph_format.page_break_before = True
            elif level == 2:
                # subtitle right under the title (before any chapter) -> centered subtitle
                if not any(lines[k].strip().startswith("# ") for k in range(i)) or not first_h1_seen:
                    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = p.add_run(strip_inline(text)); r.italic = True; r.font.size = Pt(14)
                else:
                    doc.add_heading(strip_inline(text), level=2)
            else:
                doc.add_heading(strip_inline(text), level=min(level, 4))
            i += 1
            continue

        # blockquote
        if line.startswith(">"):
            quote_lines = []
            while i < n and lines[i].strip().startswith(">"):
                quote_lines.append(lines[i].strip().lstrip(">").strip())
                i += 1
            p = doc.add_paragraph(style="Intense Quote" if "Intense Quote" in [s.name for s in doc.styles] else None)
            if p.style is None:
                p = doc.add_paragraph()
            add_runs(p, " ".join(ql for ql in quote_lines if ql))
            continue

        # unordered list
        if re.match(r'[-*+]\s+', line):
            add_runs(doc.add_paragraph(style="List Bullet"), re.sub(r'^[-*+]\s+', '', line))
            i += 1
            continue

        # ordered list
        if re.match(r'\d+\.\s+', line):
            add_runs(doc.add_paragraph(style="List Number"), re.sub(r'^\d+\.\s+', '', line))
            i += 1
            continue

        # normal paragraph
        add_runs(doc.add_paragraph(), line)
        i += 1

    doc.save(str(out_path))


def _add_table(doc, block: list) -> None:
    rows = []
    for ln in block:
        if re.fullmatch(r'\|[\s:\-|]+\|', ln):  # separator row
            continue
        cells = [c.strip() for c in ln.strip().strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=ncols)
    try:
        table.style = "Light Grid Accent 1"
    except Exception:
        table.style = "Table Grid"
    for ridx, r in enumerate(rows):
        cells = table.add_row().cells
        for cidx in range(ncols):
            txt = r[cidx] if cidx < len(r) else ""
            para = cells[cidx].paragraphs[0]
            add_runs(para, txt)
            if ridx == 0:
                for run in para.runs:
                    run.bold = True


# ----------------------------------------------------------------------
# PDF (pure-python engine, no system deps)
# ----------------------------------------------------------------------
def build_pdf(md_text: str, out_path: Path) -> None:
    import markdown as md_lib
    from xhtml2pdf import pisa

    css = """
    @page { size: A5; margin: 2cm 1.8cm; }
    body { font-family: Georgia, serif; font-size: 11pt; line-height: 1.5; }
    h1 { font-size: 20pt; page-break-before: always; }
    h1.first { page-break-before: avoid; }
    h2 { font-size: 14pt; } h3 { font-size: 12pt; } h4 { font-size: 11pt; }
    table { border-collapse: collapse; width: 100%; font-size: 9.5pt; }
    th,td { border: 1px solid #999; padding: 4px 6px; text-align: left; vertical-align: top; }
    th { background: #f0f0f0; }
    blockquote { color: #444; border-left: 3px solid #bbb; padding: 0.2em 0.9em; }
    hr { border: 0; border-top: 1px solid #ccc; margin: 1.2em 0; }
    """
    body = md_lib.markdown(md_text, extensions=["tables", "extra", "sane_lists", "nl2br"])
    first = [True]
    def mark(m):
        if first[0]:
            first[0] = False
            return '<h1 class="first"' + m.group(1) + '>'
        return '<h1' + m.group(1) + '>'
    body = re.sub(r'<h1([^>]*)>', mark, body)
    html = f"<html><head><meta charset='utf-8'><style>{css}</style></head><body>{body}</body></html>"
    with out_path.open("wb") as fh:
        pisa.CreatePDF(html, dest=fh, encoding="utf-8")


def main() -> None:
    src = Path(sys.argv[1] if len(sys.argv) > 1 else "output/full_second_draft.md")
    if not src.exists():
        raise SystemExit(f"Source not found: {src}")

    md_text = src.read_text(encoding="utf-8")

    docx_path = src.with_suffix(".docx")
    build_docx(md_text, docx_path)
    print(f"DOCX -> {docx_path}  ({docx_path.stat().st_size // 1024} KB)")

    pdf_path = src.with_suffix(".pdf")
    build_pdf(md_text, pdf_path)
    print(f"PDF  -> {pdf_path}  ({pdf_path.stat().st_size // 1024} KB)")

    print(f"Source .md left untouched: {src}")


if __name__ == "__main__":
    main()
