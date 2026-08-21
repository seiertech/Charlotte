#!/usr/bin/env python3
"""
make_report_docx.py

Render output/editorial_change_report.md into a cleanly-formatted DOCX
(output/editorial_change_report.docx), reusing the typography + markdown
rendering helpers from make_docx.py so the report matches the book's styling
(Georgia body, proper spacing, real Word tables, no blank-paragraph gaps).

Run: python3 orchestrator/make_report_docx.py
"""
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Reuse the fixed rendering + styling engine
import make_docx as md

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)
SRC = os.path.join(ROOT, "output", "editorial_change_report.md")
OUT = os.path.join(ROOT, "output", "editorial_change_report.docx")

REPORT_TITLE = "Editorial Change Report"
REPORT_SUBTITLE = "THE SYSTEM OF YOU — Draft 3 to Draft 4"


def main():
    with open(SRC, "r", encoding="utf-8") as fh:
        text = fh.read()

    # Drop the leading H1 + subtitle italic line from the body; we render a
    # proper title page instead.
    lines = text.split("\n")
    # find first H2 (## ) — everything from there is the real body
    start = 0
    for i, ln in enumerate(lines):
        if ln.startswith("## "):
            start = i
            break
    body = "\n".join(lines[start:])

    document = Document()
    md.configure_document_styles(document)

    # ---- title page ----
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_before = Pt(120)
    title = document.add_paragraph(style="Title")
    title.add_run(REPORT_TITLE)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        sub = document.add_paragraph(style="Subtitle")
    except KeyError:
        sub = document.add_paragraph(style="Heading 2")
    r = sub.add_run(REPORT_SUBTITLE)
    r.italic = True
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_before = Pt(12)
    document.add_page_break()

    # ---- body (first_chapter=True so no leading page break before first H1;
    #      but our body starts at H2, so page breaks only fire on any H1) ----
    md.render_markdown_to_docx(document, body, first_chapter=True)

    document.save(OUT)

    # sanity
    d = Document(OUT)
    paras = d.paragraphs
    print("=" * 52)
    print("REPORT DOCX BUILT")
    print("=" * 52)
    print("File   :", OUT)
    print("Size   :", os.path.getsize(OUT), "bytes")
    print("Paras  :", len(paras))
    print("Tables :", len(d.tables))
    print("H2 sections:", sum(1 for p in paras if p.style.name == "Heading 2"))
    print("H3 sections:", sum(1 for p in paras if p.style.name == "Heading 3"))


if __name__ == "__main__":
    main()
